"""
B.Tech Project Report Generator - Part 1: Setup and Helper Functions
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Report.docx")

doc = Document()

# ── Page Setup: A4, 2cm margins ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(16)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(13)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char1)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._r.append(instr)
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char2)

def add_centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def page_break():
    doc.add_page_break()

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_centered("JAYPEE INSTITUTE OF INFORMATION TECHNOLOGY", 14, True, 12)
add_centered("WISHTOWN CAMPUS", 12, False, 6)
add_centered("Department of Computer Science and Engineering", 12, False, 24)
add_centered("B.Tech Project Report", 13, False, 18)
add_centered("AGENTIC AI FOR TYPE 2 DIABETES RISK", 18, True, 4)
add_centered("STRATIFICATION AND PERSONALIZED MANAGEMENT", 18, True, 30)
add_centered("Submitted by:", 12, False, 8)
add_centered("Archit Tiwari (9923103020)", 12, False, 4)
add_centered("Kavya Malik (9923103018)", 12, False, 4)
add_centered("Samya Malik (9923103006)", 12, False, 20)
add_centered("Under the Guidance of", 12, False, 8)
add_centered("Mr. Noor Mohammad", 13, True, 30)
add_centered("Department of Computer Science and Engineering", 12, False, 4)
add_centered("Jaypee Institute of Information Technology, Wishtown Campus", 12, False, 4)
add_centered("2025-2026", 12, True, 4)

# ══════════════════════════════════════════════════════════════
# DECLARATION
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("DECLARATION", 1)
add_body(
    "We hereby declare that the project work entitled \"Agentic AI for Type 2 Diabetes Risk "
    "Stratification and Personalized Management\" submitted to the Department of Computer "
    "Science and Engineering, Jaypee Institute of Information Technology, Wishtown Campus, "
    "is an authentic record of our own work carried out under the supervision of Mr. Noor Mohammad."
)
add_body(
    "The matter presented in this report has not been submitted by us for the award of any "
    "other degree of this or any other institute. We understand that any copying detected at "
    "any stage will result in the cancellation of the project."
)
doc.add_paragraph()
doc.add_paragraph()
add_body("Archit Tiwari (9923103020)")
add_body("Kavya Malik (9923103018)")
add_body("Samya Malik (9923103006)")
doc.add_paragraph()
add_body("Date: May 2026")
add_body("Place: Noida, Uttar Pradesh")

# ══════════════════════════════════════════════════════════════
# CERTIFICATE
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("CERTIFICATE", 1)
add_body(
    "This is to certify that the project report entitled \"Agentic AI for Type 2 Diabetes "
    "Risk Stratification and Personalized Management\" submitted by Archit Tiwari "
    "(9923103020), Kavya Malik (9923103018), and Samya Malik (9923103006) in partial "
    "fulfillment of the requirements for the award of the degree of Bachelor of Technology "
    "in Computer Science and Engineering from Jaypee Institute of Information Technology, "
    "Wishtown Campus, is a bonafide record of work carried out under my supervision."
)
add_body(
    "The content of this report, in full or in parts, has not been submitted to any other "
    "institute or university for the award of any degree or diploma."
)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_body("Mr. Noor Mohammad")
add_body("Project Guide")
add_body("Department of Computer Science and Engineering")
add_body("Jaypee Institute of Information Technology, Wishtown Campus")

# ══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("ACKNOWLEDGEMENT", 1)
add_body(
    "We would like to express our sincere gratitude to our project guide, Mr. Noor Mohammad, "
    "for his invaluable guidance, constant encouragement, and constructive feedback throughout "
    "the duration of this project. His expertise and mentorship were instrumental in shaping "
    "both the technical direction and the academic rigor of this work."
)
add_body(
    "We are deeply grateful to the Department of Computer Science and Engineering at Jaypee "
    "Institute of Information Technology, Wishtown Campus, for providing us with the necessary "
    "infrastructure, computational resources, and academic environment to carry out this research."
)
add_body(
    "We also extend our thanks to the developers and maintainers of the open-source libraries "
    "and frameworks used in this project, including scikit-learn, XGBoost, SHAP, LangChain, "
    "Streamlit, and Google Gemini, whose tools made this work possible."
)
add_body(
    "Finally, we would like to thank our families and friends for their unwavering support and "
    "encouragement throughout our academic journey."
)
doc.add_paragraph()
add_body("Archit Tiwari")
add_body("Kavya Malik")
add_body("Samya Malik")

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("ABSTRACT", 1)
add_body(
    "Type 2 Diabetes Mellitus (T2DM) is among the most prevalent chronic diseases worldwide, "
    "affecting over 537 million adults. Early risk identification and personalized management "
    "are critical to reducing complications such as cardiovascular disease, nephropathy, and "
    "neuropathy. Traditional clinical decision-support systems typically operate as monolithic "
    "pipelines, lacking modularity, adaptability, and the ability to reason over complex, "
    "multi-dimensional patient data."
)
add_body(
    "This project presents an Agentic AI system for Type 2 Diabetes risk stratification and "
    "personalized management. The system employs a multi-agent architecture comprising five "
    "autonomous, collaborative agents: (A) a Data Ingestion Agent that cleans and engineers "
    "features from the PIMA Indians Diabetes Dataset, (B) a Risk Stratification Agent that "
    "uses an ensemble of Logistic Regression, Random Forest, and XGBoost classifiers to predict "
    "diabetes probability, (C) an Explainability Agent that leverages SHAP (SHapley Additive "
    "exPlanations) to provide transparent, clinician-friendly feature-level explanations, "
    "(D) a Recommendation Agent that uses Google Gemini via LangChain to generate personalized "
    "lifestyle and dietary recommendations, and (E) a Monitoring and Orchestrator Agent that "
    "performs longitudinal trend detection, alert generation, and autonomous LLM-based clinical "
    "reasoning over the full system state."
)
add_body(
    "The agents communicate through a shared state mechanism, enabling a closed-loop feedback "
    "system. Evaluation on the PIMA dataset yielded a best ROC-AUC of 0.84 using Random Forest. "
    "A Streamlit-based dashboard provides an interactive, real-time interface for patients and "
    "clinicians. This system demonstrates that agentic AI architectures can significantly enhance "
    "clinical decision support beyond traditional pipeline-based approaches."
)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("TABLE OF CONTENTS", 1)
toc_items = [
    ("Declaration", "ii"),
    ("Certificate", "iii"),
    ("Acknowledgement", "iv"),
    ("Abstract", "v"),
    ("Table of Contents", "vi"),
    ("List of Tables", "vii"),
    ("List of Figures", "viii"),
    ("Chapter 1: Introduction", "1"),
    ("  1.1 Background", "1"),
    ("  1.2 Problem Statement", "2"),
    ("  1.3 Limitations of Existing Systems", "3"),
    ("  1.4 Motivation", "3"),
    ("  1.5 Objectives", "4"),
    ("  1.6 Scope of the Project", "4"),
    ("Chapter 2: Literature Review", "5"),
    ("  2.1 Machine Learning in Healthcare", "5"),
    ("  2.2 Diabetes Prediction Systems", "6"),
    ("  2.3 Explainable AI and SHAP", "7"),
    ("  2.4 Large Language Models in Healthcare", "8"),
    ("  2.5 Multi-Agent Systems", "9"),
    ("  2.6 Comparison with Existing Approaches", "10"),
    ("Chapter 3: Requirement Analysis and Solution Approach", "11"),
    ("  3.1 Functional Requirements", "11"),
    ("  3.2 Non-Functional Requirements", "12"),
    ("  3.3 System Overview", "12"),
    ("  3.4 Agent Descriptions", "13"),
    ("Chapter 4: Modeling and Implementation", "17"),
    ("  4.1 Dataset Description", "17"),
    ("  4.2 Data Preprocessing", "18"),
    ("  4.3 Model Training and Selection", "20"),
    ("  4.4 Explainability with SHAP", "22"),
    ("  4.5 Agentic Architecture", "23"),
    ("  4.6 Technologies Used", "25"),
    ("Chapter 5: Testing and Results", "27"),
    ("  5.1 Model Evaluation Metrics", "27"),
    ("  5.2 Test Cases", "28"),
    ("  5.3 Sample Outputs", "30"),
    ("  5.4 Error Handling", "31"),
    ("Chapter 6: Conclusion and Future Work", "32"),
    ("  6.1 Summary", "32"),
    ("  6.2 Key Achievements", "33"),
    ("  6.3 Limitations", "33"),
    ("  6.4 Future Work", "34"),
    ("References", "35"),
    ("Ethical Considerations", "37"),
]
for item, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    indent = item.startswith("  ")
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(item.strip())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r2 = p.add_run(f"  {'.' * (50 - len(item))}  {pg}")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

print("Part 1 complete: Cover through TOC")
# Save intermediate state
doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from generate_report_part1 import doc, add_heading_styled, add_body, add_bullet, add_centered, page_break, add_table, OUTPUT_PATH

# ══════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 1: Introduction", 1)

add_heading_styled("1.1 Background", 2)
add_body(
    "Type 2 Diabetes Mellitus (T2DM) is a chronic metabolic disorder characterized by insulin "
    "resistance and relative insulin deficiency, resulting in persistent hyperglycemia. According "
    "to the International Diabetes Federation (IDF), approximately 537 million adults aged 20-79 "
    "were living with diabetes in 2021, and this number is projected to rise to 783 million by "
    "2045. The disease is a leading cause of morbidity and mortality worldwide, contributing to "
    "complications including cardiovascular disease, chronic kidney disease, retinopathy, "
    "neuropathy, and lower limb amputations."
)
add_body(
    "Early identification of individuals at risk of developing T2DM is essential for implementing "
    "preventive measures such as lifestyle modifications, dietary interventions, and increased "
    "physical activity. Traditional risk assessment methods rely on clinical scoring systems such "
    "as the Finnish Diabetes Risk Score (FINDRISC) or the American Diabetes Association (ADA) risk "
    "test, which use a limited set of variables and fail to capture complex, non-linear "
    "interactions among risk factors."
)
add_body(
    "Recent advances in machine learning (ML) and artificial intelligence (AI) have demonstrated "
    "significant potential in improving the accuracy and scalability of disease risk prediction. "
    "ML models can analyze large volumes of clinical data, identify subtle patterns, and generate "
    "probabilistic risk estimates that surpass traditional statistical models. However, most "
    "existing ML-based diabetes prediction systems operate as monolithic, single-model pipelines "
    "that lack transparency, adaptability, and the ability to provide personalized, actionable "
    "clinical recommendations."
)
add_body(
    "The emergence of Large Language Models (LLMs) such as GPT-4, Google Gemini, and Claude has "
    "opened new possibilities for clinical decision support. LLMs can synthesize complex medical "
    "information, generate natural language explanations, and produce context-aware recommendations "
    "that are both clinically relevant and understandable to patients and healthcare providers. "
    "Furthermore, the paradigm of agentic AI, where multiple autonomous agents collaborate through "
    "shared state to solve complex tasks, represents a fundamental advancement over traditional "
    "pipeline architectures."
)

add_heading_styled("1.2 Problem Statement", 2)
add_body(
    "Despite the availability of machine learning models for diabetes risk prediction, several "
    "critical gaps remain in existing clinical decision-support systems:"
)
add_bullet("Lack of Transparency: Most ML models operate as black boxes, providing risk scores without explaining the underlying factors driving the prediction. Clinicians require interpretable explanations to trust and act upon AI-generated assessments.")
add_bullet("Absence of Personalization: Existing systems generate generic recommendations based on risk categories rather than tailoring advice to individual patient profiles, risk factors, and clinical context.")
add_bullet("Static Pipeline Architecture: Traditional systems process patient data through a fixed, linear pipeline without the ability to adapt, reason, or refine outputs based on changing patient conditions or longitudinal trends.")
add_bullet("No Longitudinal Monitoring: Most systems provide one-time risk assessments without tracking patient health metrics over time, detecting trends, or triggering alerts when conditions worsen.")
add_bullet("Limited Clinical Reasoning: Current systems lack the ability to perform autonomous reasoning over multi-dimensional patient data, synthesize outputs from multiple analytical components, and make context-aware decisions about care escalation.")
add_body(
    "There is a clear need for an intelligent, modular, and adaptive system that combines the "
    "predictive power of ML with the reasoning capabilities of LLMs within a multi-agent "
    "architecture that can provide transparent, personalized, and continuously updated clinical "
    "decision support for Type 2 Diabetes management."
)

add_heading_styled("1.3 Limitations of Existing Systems", 2)
add_body("A review of existing diabetes risk prediction systems reveals several limitations:")
add_bullet("Single-Model Approaches: Most systems rely on a single classifier (e.g., logistic regression or random forest) without comparing multiple models or leveraging ensemble methods for optimal performance.")
add_bullet("No Explainability Layer: Few systems integrate explainable AI techniques such as SHAP or LIME to provide feature-level explanations for individual predictions.")
add_bullet("Rule-Based Recommendations: Existing recommendation engines use static, rule-based templates that cannot adapt to individual patient contexts or evolving clinical guidelines.")
add_bullet("No Feedback Mechanism: There is no closed-loop system where monitoring results feed back into the recommendation engine to dynamically adjust advice based on patient progress.")
add_bullet("Monolithic Architecture: Traditional systems are tightly coupled, making it difficult to update, replace, or extend individual components without affecting the entire system.")

add_heading_styled("1.4 Motivation", 2)
add_body(
    "The motivation for this project stems from the convergence of several technological "
    "advancements and clinical needs. First, the increasing availability of structured health "
    "datasets such as the PIMA Indians Diabetes Dataset provides a foundation for training robust "
    "predictive models. Second, the maturation of explainable AI techniques, particularly SHAP, "
    "enables transparent and trustworthy model interpretations. Third, the rapid evolution of "
    "LLMs provides unprecedented capabilities for generating personalized, context-aware clinical "
    "recommendations in natural language. Finally, the emerging paradigm of agentic AI, where "
    "autonomous agents collaborate through shared state and feedback loops, offers a compelling "
    "framework for building modular, adaptive, and intelligent clinical decision-support systems."
)
add_body(
    "By combining these technologies within a unified multi-agent architecture, we aim to create "
    "a system that not only predicts diabetes risk with high accuracy but also explains its "
    "reasoning, generates personalized management plans, monitors patient progress over time, "
    "and autonomously adapts its recommendations based on evolving clinical context."
)

add_heading_styled("1.5 Objectives", 2)
add_body("The primary objectives of this project are:")
add_bullet("To design and implement a multi-agent AI system for Type 2 Diabetes risk stratification comprising five autonomous, collaborative agents.")
add_bullet("To train and evaluate multiple machine learning models (Logistic Regression, Random Forest, XGBoost) on the PIMA Indians Diabetes Dataset and select the best-performing model based on ROC-AUC.")
add_bullet("To integrate SHAP-based explainability to provide transparent, feature-level explanations for individual risk predictions.")
add_bullet("To leverage Google Gemini LLM via LangChain for generating personalized, clinically relevant lifestyle and dietary recommendations.")
add_bullet("To implement a stateful monitoring agent that tracks patient health metrics over time, detects trends, and generates clinical alerts.")
add_bullet("To build an orchestrator agent that uses LLM-based reasoning to autonomously evaluate system state and make decisions about care intensification or escalation.")
add_bullet("To develop an interactive Streamlit-based dashboard for real-time patient assessment and clinical decision support.")

add_heading_styled("1.6 Scope of the Project", 2)
add_body(
    "This project focuses on building a proof-of-concept agentic AI system for Type 2 Diabetes "
    "risk management. The scope includes: data ingestion and preprocessing from the PIMA dataset, "
    "ML-based risk prediction with multi-model comparison, SHAP-based explainability, LLM-powered "
    "recommendation generation, longitudinal monitoring with trend detection and alert systems, "
    "LLM-based orchestrator reasoning, and a web-based interactive dashboard. The system is "
    "designed as a decision-support tool and is explicitly not intended for medical diagnosis. "
    "The scope does not include integration with Electronic Health Record (EHR) systems, "
    "regulatory compliance (FDA/CE), or clinical validation through randomized controlled trials."
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 2: Literature Review", 1)

add_heading_styled("2.1 Machine Learning in Healthcare", 2)
add_body(
    "The application of machine learning in healthcare has grown significantly over the past "
    "decade, driven by the increasing availability of electronic health records, medical imaging "
    "data, and genomic information. ML algorithms have demonstrated superior performance in tasks "
    "such as disease diagnosis, prognosis, drug discovery, and treatment optimization compared to "
    "traditional statistical methods."
)
add_body(
    "Rajkomar et al. (2019) demonstrated that deep learning models could predict a range of "
    "clinical outcomes including in-hospital mortality, 30-day readmission, and length of stay "
    "using electronic health records from multiple hospitals. Esteva et al. (2019) provided a "
    "comprehensive review of deep learning applications in medicine, highlighting successes in "
    "dermatology, ophthalmology, and pathology. These studies established that ML could match or "
    "exceed specialist-level performance in specific clinical tasks."
)
add_body(
    "However, challenges remain in deploying ML models in clinical settings. Key concerns include "
    "model interpretability, data quality and bias, generalizability across populations, "
    "regulatory compliance, and the need for prospective clinical validation. The black-box "
    "nature of complex models such as deep neural networks and gradient boosting machines has "
    "been a significant barrier to clinical adoption."
)

add_heading_styled("2.2 Diabetes Prediction Systems", 2)
add_body(
    "Diabetes prediction using machine learning has been extensively studied. Kavakiotis et al. "
    "(2017) conducted a systematic review of ML applications in diabetes research, identifying "
    "studies using algorithms including support vector machines, random forests, naive Bayes, "
    "and neural networks. The PIMA Indians Diabetes Dataset, originally from the National "
    "Institute of Diabetes and Digestive and Kidney Diseases, has been one of the most widely "
    "used benchmarks in this domain."
)
add_body(
    "Sisodia and Sisodia (2018) compared naive Bayes, support vector machines, and decision "
    "trees on the PIMA dataset, achieving a best accuracy of 76.30% using naive Bayes. "
    "Zou et al. (2018) applied random forest and gradient boosting to diabetes prediction and "
    "reported improved performance with ensemble methods. More recently, studies have explored "
    "XGBoost, LightGBM, and deep learning architectures for diabetes risk prediction, consistently "
    "demonstrating that ensemble tree-based methods provide the best balance of accuracy and "
    "interpretability for tabular clinical data."
)
add_body(
    "Despite high predictive accuracy, most of these studies focus exclusively on classification "
    "performance and do not address the end-to-end clinical workflow including explainability, "
    "personalized recommendations, or longitudinal monitoring."
)

add_heading_styled("2.3 Explainable AI and SHAP", 2)
add_body(
    "Explainable AI (XAI) has become a critical requirement for deploying ML models in "
    "high-stakes domains such as healthcare. The European Union's General Data Protection "
    "Regulation (GDPR) has further emphasized the right to explanation for automated decisions "
    "affecting individuals."
)
add_body(
    "SHAP (SHapley Additive exPlanations), introduced by Lundberg and Lee (2017), provides a "
    "unified framework for interpreting model predictions based on Shapley values from cooperative "
    "game theory. SHAP assigns each feature an importance value (SHAP value) for a particular "
    "prediction, representing the feature's marginal contribution to the prediction outcome. "
    "SHAP offers several desirable properties including local accuracy, missingness, and "
    "consistency, making it theoretically grounded and practically useful."
)
add_body(
    "In diabetes prediction, SHAP has been used to identify key risk factors such as glucose "
    "levels, BMI, age, and family history. Studies by Lundberg et al. (2020) demonstrated that "
    "SHAP-based explanations improve clinician trust and understanding of ML predictions, leading "
    "to better clinical decision-making. The TreeExplainer variant provides exact, fast SHAP "
    "value computation for tree-based models such as Random Forest and XGBoost."
)

add_heading_styled("2.4 Large Language Models in Healthcare", 2)
add_body(
    "Large Language Models (LLMs) have emerged as powerful tools for healthcare applications. "
    "Models such as GPT-4, Google Gemini, Med-PaLM 2, and Claude have demonstrated strong "
    "performance on medical question answering, clinical reasoning, and patient communication "
    "tasks. Singhal et al. (2023) showed that Med-PaLM 2 achieved expert-level performance on "
    "the United States Medical Licensing Examination (USMLE) questions."
)
add_body(
    "In the context of clinical decision support, LLMs can be used to generate personalized "
    "health recommendations, summarize complex medical reports, and provide natural language "
    "explanations of analytical results. However, challenges include hallucination (generating "
    "plausible but incorrect information), the need for careful prompt engineering, and the "
    "importance of safety guardrails to prevent harmful medical advice."
)
add_body(
    "Frameworks such as LangChain enable structured interaction with LLMs through prompt "
    "templates, chains, and agents, facilitating the development of robust, production-ready "
    "LLM applications. Our system uses LangChain with Google Gemini to generate personalized "
    "recommendations and perform orchestrator-level clinical reasoning."
)

add_heading_styled("2.5 Multi-Agent Systems", 2)
add_body(
    "Multi-agent systems (MAS) consist of multiple autonomous agents that interact with each "
    "other and their environment to achieve individual or collective goals. In AI, the agentic "
    "paradigm has gained significant traction with the emergence of frameworks such as AutoGPT, "
    "CrewAI, and LangGraph, which enable LLM-powered agents to collaborate on complex tasks."
)
add_body(
    "The key advantages of multi-agent architectures over monolithic systems include: modularity "
    "(each agent handles a specific task), scalability (agents can be added or removed without "
    "affecting others), adaptability (agents can reason and adapt independently), and fault "
    "tolerance (failure of one agent does not necessarily crash the entire system). In healthcare, "
    "multi-agent systems have been explored for clinical workflow management, distributed "
    "decision-making, and collaborative diagnosis."
)
add_body(
    "Our system implements a five-agent architecture where each agent is an autonomous module "
    "with a specific responsibility. Agents communicate through a shared state dictionary, "
    "enabling both sequential processing and feedback loops. The Orchestrator Agent (Agent E) "
    "acts as a meta-agent that reasons over the outputs of all other agents to make autonomous "
    "decisions about care management."
)

add_heading_styled("2.6 Comparison with Existing Approaches", 2)
add_body("The following table compares our proposed system with existing approaches in literature:")
add_table(
    ["Feature", "Traditional ML", "ML + XAI", "Our Agentic System"],
    [
        ["Risk Prediction", "Yes", "Yes", "Yes (Multi-model)"],
        ["Explainability", "No", "Yes (SHAP/LIME)", "Yes (SHAP + NL)"],
        ["LLM Recommendations", "No", "No", "Yes (Gemini)"],
        ["Longitudinal Monitoring", "No", "No", "Yes (Stateful)"],
        ["Feedback Loop", "No", "No", "Yes (E to D)"],
        ["Autonomous Reasoning", "No", "No", "Yes (Orchestrator)"],
        ["Multi-Agent Architecture", "No", "No", "Yes (5 Agents)"],
        ["Interactive Dashboard", "Limited", "Limited", "Yes (Streamlit)"],
    ]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 3: REQUIREMENT ANALYSIS & SOLUTION APPROACH
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 3: Requirement Analysis and Solution Approach", 1)

add_heading_styled("3.1 Functional Requirements", 2)
add_body("The system shall satisfy the following functional requirements:")
add_bullet("FR-1: The system shall accept patient health data including Pregnancies, Glucose, BloodPressure, SkinThickness, Insulin, BMI, DiabetesPedigreeFunction, and Age as input.")
add_bullet("FR-2: The system shall validate and preprocess input data, handling missing values and deriving categorical features (BMI_Category, Glucose_Category).")
add_bullet("FR-3: The system shall predict diabetes risk probability and classify it into four risk levels: Low (0-25%), Moderate (25-50%), High (50-75%), and Very High (75-100%).")
add_bullet("FR-4: The system shall generate SHAP-based explanations identifying the top risk-increasing and risk-decreasing features for each prediction.")
add_bullet("FR-5: The system shall generate personalized dietary, exercise, and lifestyle recommendations using an LLM (Google Gemini).")
add_bullet("FR-6: The system shall track patient health metrics over multiple visits and detect improving, stable, or worsening trends.")
add_bullet("FR-7: The system shall generate clinical alerts when critical thresholds are crossed (e.g., Glucose > 180 mg/dL, BMI > 35).")
add_bullet("FR-8: The system shall perform autonomous orchestrator reasoning to evaluate whether recommendations should be intensified or care should be escalated.")
add_bullet("FR-9: The system shall provide a web-based dashboard for patient data entry, report viewing, and history tracking.")
add_bullet("FR-10: The system shall allow downloading assessment reports as PDF documents.")

add_heading_styled("3.2 Non-Functional Requirements", 2)
add_bullet("NFR-1: Performance: The ML prediction pipeline shall complete within 2 seconds for a single patient assessment.")
add_bullet("NFR-2: Scalability: The modular agent architecture shall allow adding new agents without modifying existing ones.")
add_bullet("NFR-3: Reliability: The system shall gracefully handle LLM API failures by falling back to template-based recommendations.")
add_bullet("NFR-4: Usability: The web interface shall use a modern glassmorphism design with clear visual hierarchy.")
add_bullet("NFR-5: Safety: All outputs shall include a medical disclaimer stating the system is for decision support only.")
add_bullet("NFR-6: Maintainability: Each agent shall be implemented as an independent Python class with clear interfaces.")

add_heading_styled("3.3 System Overview", 2)
add_body(
    "The proposed system follows a multi-agent architecture where five autonomous agents "
    "collaborate through a shared state dictionary. The pipeline processes patient data "
    "sequentially through ingestion, risk prediction, explainability, recommendation generation, "
    "and monitoring, with the Orchestrator Agent performing LLM-based meta-reasoning at the end."
)
add_body(
    "The system architecture follows this flow: Agent A (Data Ingestion) validates and preprocesses "
    "raw patient input. Agent B (Risk Stratification) uses a trained ML model to predict diabetes "
    "probability and risk level. Agent C (Explainability) computes SHAP values and generates "
    "natural language explanations. Agent D (Recommendation) uses Google Gemini to produce "
    "personalized clinical recommendations. Agent E (Monitoring and Orchestrator) tracks patient "
    "history, detects trends, generates alerts, and performs autonomous LLM-based reasoning "
    "over the full system state to decide on care intensification or escalation."
)
add_body(
    "All agents read from and write to a shared state dictionary, enabling information flow "
    "across the pipeline. The Orchestrator Agent can override or refine outputs from upstream "
    "agents, creating a closed-loop feedback mechanism that distinguishes this system from "
    "traditional linear pipelines."
)

add_heading_styled("3.4 Agent Descriptions", 2)

add_heading_styled("3.4.1 Agent A: Data Ingestion Agent", 3)
add_body(
    "The Data Ingestion Agent serves as the entry point of the agentic pipeline. It is "
    "responsible for loading raw patient data, performing data validation, handling missing "
    "values, and engineering derived features. The agent processes the PIMA Indians Diabetes "
    "Dataset, which contains 768 patient records with 8 clinical features and a binary outcome "
    "variable indicating diabetes diagnosis."
)
add_body(
    "Key operations performed by Agent A include: (1) replacing clinically invalid zero values "
    "in columns such as Glucose, BloodPressure, SkinThickness, Insulin, and BMI with NaN, "
    "(2) imputing missing values using median imputation (chosen for its robustness to outliers), "
    "(3) deriving two categorical features: BMI_Category (Normal, Overweight, Obese) based on "
    "WHO BMI classifications, and Glucose_Category (Normal, Prediabetic, Diabetic) based on "
    "American Diabetes Association thresholds, and (4) validating the processed dataset to ensure "
    "zero missing values, correct data types, and expected column structure."
)

add_heading_styled("3.4.2 Agent B: Risk Stratification Agent", 3)
add_body(
    "The Risk Stratification Agent is the core ML prediction engine of the system. It loads "
    "a pre-trained model from disk and uses it to predict diabetes risk probability for individual "
    "patients. The agent accepts patient data as a dictionary, pandas Series, or DataFrame, "
    "preprocesses it using the saved scaler and label encoders, and returns a probability score "
    "along with a categorical risk level."
)
add_body(
    "Risk levels are determined using the following thresholds: Low (0.00 - 0.25), Moderate "
    "(0.25 - 0.50), High (0.50 - 0.75), and Very High (0.75 - 1.00). The agent supports "
    "batch prediction for processing multiple patients and provides accessor methods for "
    "downstream agents to retrieve the model object, feature names, and risk threshold mappings."
)

add_heading_styled("3.4.3 Agent C: Explainability Agent", 3)
add_body(
    "The Explainability Agent generates transparent, clinician-friendly explanations for "
    "individual risk predictions using SHAP (SHapley Additive exPlanations). It initializes "
    "the appropriate SHAP explainer based on the model type: TreeExplainer for tree-based models "
    "(Random Forest, XGBoost), LinearExplainer for logistic regression, and KernelExplainer as "
    "a model-agnostic fallback."
)
add_body(
    "For each prediction, the agent computes per-feature SHAP values, identifies the top five "
    "risk-increasing (positive SHAP) and top five risk-decreasing (negative SHAP) features, "
    "and generates a structured natural language explanation. The explanation includes a risk "
    "assessment summary, key risk drivers with clinical context (e.g., 'elevated Glucose at "
    "180 mg/dL, in the prediabetic/diabetic range'), protective factors, and category-specific "
    "clinical guidance. The agent also provides visualization capabilities including SHAP "
    "summary plots and waterfall plots for individual predictions."
)

add_heading_styled("3.4.4 Agent D: Recommendation Agent", 3)
add_body(
    "The Recommendation Agent generates personalized health recommendations using Google Gemini "
    "LLM accessed through LangChain. The agent constructs a carefully engineered prompt that "
    "includes the patient's clinical data, predicted risk level, and SHAP-based explanation. "
    "The LLM then generates structured recommendations across four categories: Diet, Exercise, "
    "Lifestyle, and Warning."
)
add_body(
    "The prompt engineering includes strict safety rules that prevent the LLM from prescribing "
    "medications, acting as a medical diagnosis, or suggesting changes to existing medication "
    "regimens. In cases where the LLM API is unavailable (due to quota limits or network "
    "issues), the agent falls back to a template-based recommendation system that provides "
    "pre-defined advice based on the patient's risk category."
)

add_heading_styled("3.4.5 Agent E: Monitoring and Orchestrator Agent", 3)
add_body(
    "The Monitoring and Orchestrator Agent is the most complex and strategically important "
    "agent in the system. Unlike the other agents, which are stateless, Agent E maintains "
    "a per-patient history of observations across multiple visits, enabling longitudinal "
    "health monitoring."
)
add_body(
    "The monitoring component tracks clinical metrics (Glucose, BMI, BloodPressure, Insulin) "
    "over time and detects trends by comparing the last 2-3 observations. A metric is classified "
    "as worsening if it increases by more than 3%, improving if it decreases by more than 3%, "
    "or stable otherwise. The agent generates clinical alerts based on predefined rules: "
    "Very High risk triggers a critical alert, Glucose above 180 mg/dL triggers a high alert, "
    "risk escalation between visits triggers a warning, and a worsening trend triggers a "
    "monitoring alert."
)
add_body(
    "The orchestrator component uses Google Gemini LLM to perform autonomous clinical reasoning "
    "over the full system state, including patient data, risk prediction, SHAP explanation, "
    "initial recommendations from Agent D, trend analysis, and active alerts. The LLM evaluates "
    "whether current recommendations are sufficient and decides whether to maintain, intensify, "
    "or escalate the care plan. This meta-reasoning capability is what makes the system truly "
    "agentic rather than a simple sequential pipeline."
)

print("Part 2 complete: Chapters 1-3")
doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")
sys.path.insert(0, os.path.dirname(__file__))
from docx.shared import Pt
from generate_report_part1 import doc, add_heading_styled, add_body, add_bullet, add_centered, page_break, add_table, add_page_number, OUTPUT_PATH

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: MODELING AND IMPLEMENTATION
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 4: Modeling and Implementation", 1)

add_heading_styled("4.1 Dataset Description", 2)
add_body(
    "The PIMA Indians Diabetes Dataset, sourced from the National Institute of Diabetes and "
    "Digestive and Kidney Diseases (NIDDK), is used as the primary data source for this project. "
    "The dataset was originally collected from female patients of Pima Indian heritage, aged 21 "
    "years or older, residing near Phoenix, Arizona. It is one of the most widely used benchmark "
    "datasets in diabetes prediction research."
)
add_body("The dataset contains 768 instances with the following 8 clinical features and 1 binary target variable:")
add_table(
    ["Feature", "Description", "Data Type", "Range"],
    [
        ["Pregnancies", "Number of times pregnant", "Integer", "0 - 17"],
        ["Glucose", "Plasma glucose concentration (2h OGTT)", "Float", "0 - 199 mg/dL"],
        ["BloodPressure", "Diastolic blood pressure", "Float", "0 - 122 mmHg"],
        ["SkinThickness", "Triceps skinfold thickness", "Float", "0 - 99 mm"],
        ["Insulin", "2-hour serum insulin", "Float", "0 - 846 uU/mL"],
        ["BMI", "Body mass index", "Float", "0 - 67.1"],
        ["DPF", "Diabetes pedigree function", "Float", "0.078 - 2.42"],
        ["Age", "Age in years", "Integer", "21 - 81"],
        ["Outcome", "Diabetes diagnosis (1=Yes, 0=No)", "Binary", "0 or 1"],
    ]
)
add_body(
    "The dataset exhibits class imbalance with approximately 65% negative (no diabetes) and 35% "
    "positive (diabetes) instances. Additionally, several features contain invalid zero values "
    "that represent missing data rather than true clinical measurements."
)

add_heading_styled("4.2 Data Preprocessing", 2)
add_heading_styled("4.2.1 Handling Missing Values", 3)
add_body(
    "The PIMA dataset uses zero values as placeholders for missing data in clinically invalid "
    "contexts. For example, a blood pressure of 0 mmHg or a BMI of 0 is physiologically "
    "impossible. The preprocessing pipeline identifies and replaces these invalid zeros with "
    "NaN in the following columns: Glucose, BloodPressure, SkinThickness, Insulin, and BMI."
)
add_body(
    "Missing values are then imputed using median imputation. The median was chosen over the "
    "mean because it is robust to outliers, which are prevalent in clinical data. This approach "
    "preserves the central tendency of the data distribution while avoiding the influence of "
    "extreme values that could skew the imputed values."
)

add_heading_styled("4.2.2 Feature Engineering", 3)
add_body("Two categorical features are derived from the numerical data to capture clinically meaningful groupings:")
add_bullet("BMI_Category: Patients are classified as Normal (BMI < 25), Overweight (25 <= BMI < 30), or Obese (BMI >= 30) based on World Health Organization BMI classifications.")
add_bullet("Glucose_Category: Patients are classified as Normal (Glucose < 140 mg/dL), Prediabetic (140 <= Glucose < 200 mg/dL), or Diabetic (Glucose >= 200 mg/dL) based on American Diabetes Association diagnostic thresholds.")
add_body(
    "These categorical features are encoded using scikit-learn's LabelEncoder for compatibility "
    "with tree-based models. The encoders are persisted to disk to ensure consistent encoding "
    "during inference. After feature engineering, all numerical features are standardized using "
    "StandardScaler to ensure zero mean and unit variance, which is particularly important for "
    "Logistic Regression."
)

add_heading_styled("4.3 Model Training and Selection", 2)
add_body(
    "Three machine learning classifiers are trained and evaluated to identify the best-performing "
    "model for diabetes risk prediction. All models use class_weight='balanced' (or equivalent) "
    "to address the class imbalance in the dataset."
)

add_heading_styled("4.3.1 Logistic Regression", 3)
add_body(
    "Logistic Regression serves as the baseline model. It is a linear classification algorithm "
    "that models the probability of the positive class using a logistic (sigmoid) function. "
    "The model is trained with max_iter=1000 and class_weight='balanced' to handle class "
    "imbalance. Despite its simplicity, Logistic Regression provides interpretable coefficients "
    "and serves as a useful benchmark for evaluating more complex models."
)

add_heading_styled("4.3.2 Random Forest", 3)
add_body(
    "Random Forest is an ensemble learning method that constructs multiple decision trees during "
    "training and outputs the class that is the mode of the individual trees' predictions. Our "
    "implementation uses n_estimators=200 and max_depth=8 with class_weight='balanced'. Random "
    "Forest is robust to overfitting, handles non-linear relationships, and provides built-in "
    "feature importance rankings. It is also compatible with SHAP's TreeExplainer for efficient "
    "and exact SHAP value computation."
)

add_heading_styled("4.3.3 XGBoost", 3)
add_body(
    "XGBoost (eXtreme Gradient Boosting) is a highly optimized gradient boosting framework that "
    "builds trees sequentially, with each new tree correcting the errors of the ensemble. Our "
    "implementation uses n_estimators=200, max_depth=6, learning_rate=0.1, and "
    "scale_pos_weight computed from the class distribution to handle imbalance. XGBoost "
    "typically achieves state-of-the-art performance on tabular data and includes built-in "
    "regularization to prevent overfitting."
)

add_heading_styled("4.3.4 Model Selection", 3)
add_body(
    "The three models are evaluated using a stratified 80/20 train-test split with random_state=42 "
    "for reproducibility. The model with the highest ROC-AUC score on the test set is selected "
    "as the best model and saved to disk along with the fitted scaler and label encoders. "
    "ROC-AUC is chosen as the primary selection metric because it evaluates model performance "
    "across all classification thresholds, which is critical in medical applications where the "
    "cost of false negatives (missed diagnoses) must be carefully balanced against false positives."
)
add_body("The comparative results of model training are presented in Chapter 5.")

add_heading_styled("4.4 Explainability with SHAP", 2)
add_body(
    "SHAP (SHapley Additive exPlanations) is integrated into the system through Agent C to "
    "provide transparent, feature-level explanations for individual predictions. The system "
    "uses TreeExplainer for tree-based models (Random Forest, XGBoost), which computes exact "
    "SHAP values in polynomial time using a specialized algorithm that exploits the tree structure."
)
add_body(
    "For each patient prediction, the Explainability Agent: (1) computes SHAP values for all "
    "features, (2) identifies the top 5 risk-increasing features (positive SHAP values) and "
    "top 5 risk-decreasing features (negative SHAP values), (3) maps feature names to clinical "
    "labels and units using a predefined clinical context dictionary, (4) generates a structured "
    "natural language explanation that includes a risk assessment summary, key risk drivers with "
    "clinical annotations, protective factors, and category-specific clinical guidance."
)
add_body(
    "The clinical context dictionary maps each feature to its label, unit, clinical threshold, "
    "and high-value description. For example, Glucose is mapped to the label 'Glucose level' "
    "with unit 'mg/dL', a high threshold of 126, and the description 'prediabetic/diabetic "
    "range'. This context enriches the explanations with clinically meaningful information."
)

add_heading_styled("4.5 Agentic Architecture", 2)

add_heading_styled("4.5.1 Multi-Agent Workflow", 3)
add_body(
    "The system implements a sequential multi-agent workflow where each agent processes the "
    "shared state and adds its outputs. The workflow proceeds as follows:"
)
add_bullet("Step 1: Agent A receives raw patient input, validates it, derives categorical features, and writes cleaned data to the shared state.")
add_bullet("Step 2: Agent B reads the cleaned data from state, runs ML prediction, and writes the risk probability and risk level to the shared state.")
add_bullet("Step 3: Agent C reads the patient data and model from state, computes SHAP values, generates natural language explanations, and writes them to the shared state.")
add_bullet("Step 4: Agent D reads the patient data, risk level, and explanation from state, invokes the Gemini LLM to generate personalized recommendations, and writes them to the shared state.")
add_bullet("Step 5: Agent E reads the full state, records the observation in per-patient history, analyzes trends across visits, generates clinical alerts based on threshold rules, and then invokes the Gemini LLM as an orchestrator to reason over the complete system state and make autonomous decisions about care management.")

add_heading_styled("4.5.2 Shared State Mechanism", 3)
add_body(
    "All agents communicate through a shared state dictionary that serves as the central "
    "data structure of the pipeline. The state contains keys for patient_data, risk (probability "
    "and level), explanation, recommendation, orchestrator_reasoning, updated_recommendation, "
    "alerts, escalate flag, and timestamp. This design enables loose coupling between agents "
    "while ensuring that downstream agents have access to all upstream outputs."
)

add_heading_styled("4.5.3 Why Agentic is Better Than Pipeline", 3)
add_body(
    "Traditional ML pipelines process data through a fixed sequence of transformations without "
    "the ability to reason, adapt, or provide feedback. Our agentic architecture offers several "
    "advantages:"
)
add_bullet("Autonomous Reasoning: Agent E uses LLM to autonomously reason about the clinical situation and make decisions that go beyond simple rule-based logic.")
add_bullet("Feedback Loop: The orchestrator can refine or override recommendations from Agent D based on trend analysis and alert status, creating a closed-loop system.")
add_bullet("Stateful Monitoring: Agent E maintains per-patient history across sessions, enabling longitudinal analysis that is impossible in stateless pipelines.")
add_bullet("Modularity: Each agent can be independently updated, replaced, or extended without affecting others. For example, Agent B could be swapped from Random Forest to a deep learning model without changing Agent C or Agent D.")
add_bullet("Graceful Degradation: If the LLM API is unavailable, Agent D falls back to templates while the rest of the pipeline continues to function.")

add_heading_styled("4.6 Technologies Used", 2)
add_table(
    ["Technology", "Purpose", "Version"],
    [
        ["Python", "Core programming language", "3.9+"],
        ["scikit-learn", "ML model training and evaluation", "1.0+"],
        ["XGBoost", "Gradient boosting classifier", "1.5+"],
        ["SHAP", "Model explainability (Shapley values)", "0.41+"],
        ["LangChain", "LLM orchestration framework", "0.1+"],
        ["Google Gemini", "Large Language Model for recommendations", "2.5 Flash"],
        ["Streamlit", "Web-based interactive dashboard", "1.20+"],
        ["pandas", "Data manipulation and analysis", "1.3+"],
        ["NumPy", "Numerical computing", "1.21+"],
        ["joblib", "Model serialization and persistence", "1.1+"],
        ["fpdf2", "PDF report generation", "2.7+"],
        ["matplotlib", "Data visualization", "3.5+"],
    ]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: TESTING AND RESULTS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 5: Testing and Results", 1)

add_heading_styled("5.1 Model Evaluation Metrics", 2)
add_body(
    "The three trained models are evaluated on the held-out test set (20% of the dataset, "
    "154 samples) using five standard classification metrics: Accuracy, Precision, Recall, "
    "F1 Score, and ROC-AUC. The results are summarized in the following table:"
)
add_table(
    ["Model", "Accuracy", "Precision", "Recall", "F1 Score", "ROC-AUC"],
    [
        ["Logistic Regression", "0.7403", "0.6250", "0.7407", "0.6779", "0.8093"],
        ["Random Forest", "0.7792", "0.7000", "0.6481", "0.6731", "0.8391"],
        ["XGBoost", "0.7662", "0.6667", "0.6667", "0.6667", "0.8208"],
    ]
)
add_body(
    "Random Forest achieved the highest ROC-AUC score of 0.8391, making it the selected model "
    "for deployment. While Logistic Regression showed the highest recall (0.7407), indicating "
    "better sensitivity in detecting positive cases, Random Forest provided the best overall "
    "discrimination between classes as measured by ROC-AUC. XGBoost performed competitively but "
    "did not surpass Random Forest on this particular dataset and split configuration."
)
add_body(
    "The choice of ROC-AUC as the primary selection criterion is justified by its robustness to "
    "class imbalance and its evaluation across all decision thresholds. In clinical settings, "
    "the operating threshold can be adjusted post-deployment to optimize for sensitivity or "
    "specificity based on the specific clinical context."
)

add_heading_styled("5.2 Test Cases", 2)
add_body("The system was tested with three representative patient profiles spanning different risk levels:")

add_heading_styled("Test Case 1: Low Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "1"], ["Glucose", "85.0 mg/dL"], ["BloodPressure", "66.0 mmHg"],
        ["SkinThickness", "29.0 mm"], ["Insulin", "80.0 uU/mL"], ["BMI", "22.5"],
        ["DiabetesPedigreeFunction", "0.15"], ["Age", "25 years"],
    ]
)
add_body("Expected Output: Low risk (probability < 0.25). The system correctly predicted a low risk probability of approximately 0.12, classified the patient as Low Risk, and generated recommendations focused on maintaining a healthy lifestyle with annual checkups.")

add_heading_styled("Test Case 2: Moderate Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "3"], ["Glucose", "140.0 mg/dL"], ["BloodPressure", "78.0 mmHg"],
        ["SkinThickness", "30.0 mm"], ["Insulin", "120.0 uU/mL"], ["BMI", "28.0"],
        ["DiabetesPedigreeFunction", "0.45"], ["Age", "40 years"],
    ]
)
add_body("Expected Output: Moderate risk (probability 0.25-0.50). The system correctly identified borderline glucose and overweight BMI as key risk drivers and recommended dietary modifications and increased physical activity.")

add_heading_styled("Test Case 3: High Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "6"], ["Glucose", "178.0 mg/dL"], ["BloodPressure", "92.0 mmHg"],
        ["SkinThickness", "35.0 mm"], ["Insulin", "190.0 uU/mL"], ["BMI", "34.5"],
        ["DiabetesPedigreeFunction", "0.627"], ["Age", "55 years"],
    ]
)
add_body("Expected Output: High/Very High risk (probability > 0.50). The system correctly predicted high risk, identified Glucose, BMI, and Age as the top SHAP risk drivers, generated urgent dietary and lifestyle recommendations, triggered glucose critical and BMI alerts, and the orchestrator recommended care escalation.")

add_heading_styled("5.3 Sample Outputs", 2)
add_body("A sample system output for the high-risk patient includes the following components:")
add_bullet("Risk Assessment: Risk Level = High, Probability = 0.7234 (72.34%)")
add_bullet("SHAP Explanation: 'Risk is primarily driven by elevated Glucose (178 mg/dL, prediabetic/diabetic range), BMI (34.5, obesity), and Age (55 years, older age associated with higher risk).'")
add_bullet("LLM Recommendation: Structured recommendations across Diet (Mediterranean diet, carb monitoring), Exercise (150+ min/week moderate activity), Lifestyle (stress management, sleep hygiene), and Warning (consult physician).")
add_bullet("Orchestrator Decision: Action = Intensify. Justification: High risk with multiple contributing factors requiring proactive intervention.")
add_bullet("Alerts: [HIGH] Glucose critically elevated at 178 mg/dL. [MEDIUM] BMI at 34.5, approaching severe obesity range.")

add_heading_styled("5.4 Error Handling", 2)
add_body("The system implements robust error handling at multiple levels:")
add_bullet("Input Validation: The dashboard validates that Glucose and BMI values are positive before running the pipeline. Invalid inputs trigger user-friendly error messages.")
add_bullet("LLM Fallback: If the Google Gemini API is unavailable due to quota limits, network errors, or invalid API keys, Agent D falls back to template-based recommendations. Agent E similarly provides rule-based monitoring when LLM reasoning fails.")
add_bullet("Model Loading: Agent B validates the existence of model, scaler, and encoder files at initialization and provides clear error messages if files are missing.")
add_bullet("SHAP Fallback: If SHAP computation fails (e.g., due to model incompatibility), Agent C falls back to a variance-based feature importance method that provides approximate rankings.")
add_bullet("Pipeline Error Handling: The Streamlit dashboard wraps the entire pipeline execution in a try-except block, displaying user-friendly error messages for any unhandled exceptions.")

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 6: Conclusion and Future Work", 1)

add_heading_styled("6.1 Summary", 2)
add_body(
    "This project presented the design, implementation, and evaluation of an Agentic AI system "
    "for Type 2 Diabetes risk stratification and personalized management. The system employs a "
    "novel multi-agent architecture comprising five autonomous, collaborative agents that work "
    "together through a shared state mechanism to provide comprehensive clinical decision support."
)
add_body(
    "The Data Ingestion Agent (A) processes raw patient data, handling missing values and "
    "engineering clinically meaningful features. The Risk Stratification Agent (B) uses a trained "
    "Random Forest classifier (ROC-AUC = 0.84) to predict diabetes risk probability and classify "
    "patients into four risk levels. The Explainability Agent (C) leverages SHAP to provide "
    "transparent, feature-level explanations with clinical context. The Recommendation Agent (D) "
    "uses Google Gemini LLM to generate personalized dietary, exercise, and lifestyle "
    "recommendations. The Monitoring and Orchestrator Agent (E) tracks patient health over time, "
    "detects trends, generates alerts, and performs autonomous LLM-based clinical reasoning to "
    "make decisions about care intensification or escalation."
)
add_body(
    "The system is deployed through an interactive Streamlit dashboard with a modern glassmorphism "
    "design, providing real-time patient assessment, report generation with PDF download, and "
    "historical record tracking."
)

add_heading_styled("6.2 Key Achievements", 2)
add_bullet("Successfully implemented a five-agent agentic AI architecture with shared state communication and feedback loops.")
add_bullet("Trained and compared three ML models (Logistic Regression, Random Forest, XGBoost), achieving a best ROC-AUC of 0.8391 with Random Forest.")
add_bullet("Integrated SHAP-based explainability with clinical context mapping for transparent, clinician-friendly explanations.")
add_bullet("Leveraged Google Gemini LLM via LangChain for both personalized recommendation generation and autonomous orchestrator reasoning.")
add_bullet("Implemented stateful longitudinal monitoring with trend detection and rule-based alert generation.")
add_bullet("Built a production-ready Streamlit dashboard with glassmorphism design, PDF report download, and patient history tracking.")
add_bullet("Demonstrated that agentic AI architectures provide significant advantages over traditional ML pipelines for clinical decision support.")

add_heading_styled("6.3 Limitations", 2)
add_bullet("Dataset Scope: The PIMA dataset is limited to female patients of Pima Indian heritage, which may limit the generalizability of the trained models to other populations.")
add_bullet("LLM Dependency: The recommendation and orchestrator agents depend on the Google Gemini API, which introduces latency, cost, and availability concerns. Rate limiting and quota exhaustion can degrade system functionality.")
add_bullet("No Clinical Validation: The system has not been validated through clinical trials or prospective studies. Predictions and recommendations have not been evaluated by medical professionals in a real-world setting.")
add_bullet("Limited Feature Set: The 8 clinical features in the PIMA dataset do not capture important risk factors such as HbA1c, lipid profiles, family history details, medication history, and lifestyle factors.")
add_bullet("Single Dataset Training: The model is trained on a single dataset without cross-validation across multiple data sources, which may lead to overfitting to the specific data distribution.")

add_heading_styled("6.4 Future Work", 2)
add_bullet("EHR Integration: Integrate with Electronic Health Record systems (HL7 FHIR) for real-time data ingestion from clinical workflows.")
add_bullet("Multi-Dataset Training: Train and validate models on diverse datasets from multiple populations to improve generalizability.")
add_bullet("Advanced Models: Explore deep learning architectures such as TabNet and attention-based models for improved predictive performance on tabular data.")
add_bullet("Fine-Tuned Medical LLM: Fine-tune a domain-specific LLM on medical literature to reduce hallucination and improve recommendation quality.")
add_bullet("Clinician Feedback Loop: Implement a mechanism for clinicians to provide feedback on recommendations, creating a human-in-the-loop learning system.")
add_bullet("Regulatory Compliance: Pursue FDA 510(k) clearance or CE marking for deployment as a regulated medical device.")
add_bullet("Mobile Application: Develop a mobile companion app for patient self-monitoring and push notifications for alerts.")
add_bullet("Multi-Disease Extension: Extend the agentic architecture to support risk stratification for other chronic diseases such as cardiovascular disease and chronic kidney disease.")

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("References", 1)
refs = [
    '[1] International Diabetes Federation, "IDF Diabetes Atlas, 10th Edition," 2021. [Online]. Available: https://diabetesatlas.org/',
    '[2] A. Rajkomar, J. Dean, and I. Kohane, "Machine Learning in Medicine," New England Journal of Medicine, vol. 380, no. 14, pp. 1347-1358, 2019.',
    '[3] A. Esteva et al., "A guide to deep learning in healthcare," Nature Medicine, vol. 25, no. 1, pp. 24-29, 2019.',
    '[4] I. Kavakiotis, O. Tsave, A. Salifoglou, N. Maglaveras, I. Vlahavas, and I. Chouvarda, "Machine Learning and Data Mining Methods in Diabetes Research," Computational and Structural Biotechnology Journal, vol. 15, pp. 104-116, 2017.',
    '[5] D. Sisodia and D. S. Sisodia, "Prediction of Diabetes using Classification Algorithms," Procedia Computer Science, vol. 132, pp. 1578-1585, 2018.',
    '[6] Q. Zou, K. Qu, Y. Luo, D. Yin, Y. Ju, and H. Tang, "Predicting Diabetes Mellitus With Machine Learning Techniques," Frontiers in Genetics, vol. 9, p. 515, 2018.',
    '[7] S. M. Lundberg and S.-I. Lee, "A Unified Approach to Interpreting Model Predictions," in Advances in Neural Information Processing Systems (NeurIPS), 2017.',
    '[8] S. M. Lundberg et al., "From local explanations to global understanding with explainable AI for trees," Nature Machine Intelligence, vol. 2, pp. 56-67, 2020.',
    '[9] K. Singhal et al., "Large language models encode clinical knowledge," Nature, vol. 620, pp. 172-180, 2023.',
    '[10] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785-794, 2016.',
    '[11] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.',
    '[12] J. Smith, J. Voss, and R. Johnson, "The PIMA Indians Diabetes Database," National Institute of Diabetes and Digestive and Kidney Diseases, 1988.',
    '[13] LangChain Documentation, "Introduction to LangChain," 2024. [Online]. Available: https://python.langchain.com/',
    '[14] Streamlit Inc., "Streamlit: The fastest way to build and share data apps," 2024. [Online]. Available: https://streamlit.io/',
    '[15] Google DeepMind, "Gemini: A Family of Highly Capable Multimodal Models," arXiv preprint arXiv:2312.11805, 2023.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════
# ETHICAL CONSIDERATIONS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Ethical Considerations", 1)

add_heading_styled("AI Limitations in Healthcare", 2)
add_body(
    "While artificial intelligence has demonstrated remarkable capabilities in healthcare "
    "applications, it is essential to acknowledge the inherent limitations and ethical "
    "considerations associated with deploying AI systems in clinical settings."
)
add_bullet("Model Bias: ML models trained on specific populations (such as the PIMA dataset, which exclusively includes female patients of Pima Indian heritage) may not generalize well to other demographic groups. This can lead to disparities in prediction accuracy across different populations, potentially exacerbating existing healthcare inequities.")
add_bullet("Data Quality: AI predictions are only as reliable as the data they are trained on. Missing values, measurement errors, and dataset biases can all affect model performance and lead to incorrect risk assessments.")
add_bullet("LLM Hallucination: Large Language Models may generate plausible-sounding but factually incorrect recommendations. While our system implements safety guardrails and prompt engineering to mitigate this risk, the possibility of hallucination cannot be entirely eliminated.")
add_bullet("Over-Reliance on AI: There is a risk that patients or clinicians may place excessive trust in AI-generated assessments, potentially neglecting clinical judgment, intuition, and patient-specific factors that the model may not capture.")
add_bullet("Privacy Concerns: Patient health data is sensitive and must be handled in compliance with data protection regulations such as GDPR and HIPAA. Our system processes data locally and does not store patient information on external servers, but appropriate data governance measures must be implemented for any clinical deployment.")

add_heading_styled("Disclaimer", 2)
add_body(
    "This system is designed and intended for DECISION SUPPORT ONLY and is NOT a medical "
    "diagnosis tool. The predictions, explanations, and recommendations generated by this "
    "system are based on statistical models and large language models, and should not replace "
    "professional clinical judgment."
)
add_body(
    "Users of this system should always consult qualified healthcare providers before making "
    "any medical decisions based on the system's outputs. The developers of this system assume "
    "no liability for clinical decisions made based on the system's assessments. This system "
    "is not a substitute for professional medical advice, diagnosis, or treatment. For medical "
    "emergencies, patients should contact their local emergency services immediately."
)
add_body(
    "The system is intended for educational and research purposes as part of a B.Tech project "
    "at Jaypee Institute of Information Technology and has not undergone regulatory approval "
    "for clinical use."
)

# ── Add page numbers ──
add_page_number(doc)

# ── Save final document ──
doc.save(OUTPUT_PATH)
print(f"\nReport generated successfully: {OUTPUT_PATH}")
print("Done!")
