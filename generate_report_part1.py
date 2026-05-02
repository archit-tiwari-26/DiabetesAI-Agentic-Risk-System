"""
B.Tech Project Report Generator - Part 1: Setup and Helper Functions
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Project_Report.docx")

doc = Document()

# ── Page Setup: A4, 2cm margins ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        hs.font.size = Pt(16)
    elif i == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(13)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char1)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._r.append(instr)
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char2)

def add_centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def page_break():
    doc.add_page_break()

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_centered("JAYPEE INSTITUTE OF INFORMATION TECHNOLOGY", 14, True, 12)
add_centered("WISHTOWN CAMPUS", 12, False, 6)
add_centered("Department of Computer Science and Engineering", 12, False, 24)
add_centered("B.Tech Project Report", 13, False, 18)
add_centered("AGENTIC AI FOR TYPE 2 DIABETES RISK", 18, True, 4)
add_centered("STRATIFICATION AND PERSONALIZED MANAGEMENT", 18, True, 30)
add_centered("Submitted by:", 12, False, 8)
add_centered("Archit Tiwari (9923103020)", 12, False, 4)
add_centered("Kavya Malik (9923103018)", 12, False, 4)
add_centered("Samya Malik (9923103006)", 12, False, 20)
add_centered("Under the Guidance of", 12, False, 8)
add_centered("Mr. Noor Mohammad", 13, True, 30)
add_centered("Department of Computer Science and Engineering", 12, False, 4)
add_centered("Jaypee Institute of Information Technology, Wishtown Campus", 12, False, 4)
add_centered("2025-2026", 12, True, 4)

# ══════════════════════════════════════════════════════════════
# DECLARATION
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("DECLARATION", 1)
add_body(
    "We hereby declare that the project work entitled \"Agentic AI for Type 2 Diabetes Risk "
    "Stratification and Personalized Management\" submitted to the Department of Computer "
    "Science and Engineering, Jaypee Institute of Information Technology, Wishtown Campus, "
    "is an authentic record of our own work carried out under the supervision of Mr. Noor Mohammad."
)
add_body(
    "The matter presented in this report has not been submitted by us for the award of any "
    "other degree of this or any other institute. We understand that any copying detected at "
    "any stage will result in the cancellation of the project."
)
doc.add_paragraph()
doc.add_paragraph()
add_body("Archit Tiwari (9923103020)")
add_body("Kavya Malik (9923103018)")
add_body("Samya Malik (9923103006)")
doc.add_paragraph()
add_body("Date: May 2026")
add_body("Place: Noida, Uttar Pradesh")

# ══════════════════════════════════════════════════════════════
# CERTIFICATE
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("CERTIFICATE", 1)
add_body(
    "This is to certify that the project report entitled \"Agentic AI for Type 2 Diabetes "
    "Risk Stratification and Personalized Management\" submitted by Archit Tiwari "
    "(9923103020), Kavya Malik (9923103018), and Samya Malik (9923103006) in partial "
    "fulfillment of the requirements for the award of the degree of Bachelor of Technology "
    "in Computer Science and Engineering from Jaypee Institute of Information Technology, "
    "Wishtown Campus, is a bonafide record of work carried out under my supervision."
)
add_body(
    "The content of this report, in full or in parts, has not been submitted to any other "
    "institute or university for the award of any degree or diploma."
)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
add_body("Mr. Noor Mohammad")
add_body("Project Guide")
add_body("Department of Computer Science and Engineering")
add_body("Jaypee Institute of Information Technology, Wishtown Campus")

# ══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("ACKNOWLEDGEMENT", 1)
add_body(
    "We would like to express our sincere gratitude to our project guide, Mr. Noor Mohammad, "
    "for his invaluable guidance, constant encouragement, and constructive feedback throughout "
    "the duration of this project. His expertise and mentorship were instrumental in shaping "
    "both the technical direction and the academic rigor of this work."
)
add_body(
    "We are deeply grateful to the Department of Computer Science and Engineering at Jaypee "
    "Institute of Information Technology, Wishtown Campus, for providing us with the necessary "
    "infrastructure, computational resources, and academic environment to carry out this research."
)
add_body(
    "We also extend our thanks to the developers and maintainers of the open-source libraries "
    "and frameworks used in this project, including scikit-learn, XGBoost, SHAP, LangChain, "
    "Streamlit, and Google Gemini, whose tools made this work possible."
)
add_body(
    "Finally, we would like to thank our families and friends for their unwavering support and "
    "encouragement throughout our academic journey."
)
doc.add_paragraph()
add_body("Archit Tiwari")
add_body("Kavya Malik")
add_body("Samya Malik")

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("ABSTRACT", 1)
add_body(
    "Type 2 Diabetes Mellitus (T2DM) is among the most prevalent chronic diseases worldwide, "
    "affecting over 537 million adults. Early risk identification and personalized management "
    "are critical to reducing complications such as cardiovascular disease, nephropathy, and "
    "neuropathy. Traditional clinical decision-support systems typically operate as monolithic "
    "pipelines, lacking modularity, adaptability, and the ability to reason over complex, "
    "multi-dimensional patient data."
)
add_body(
    "This project presents an Agentic AI system for Type 2 Diabetes risk stratification and "
    "personalized management. The system employs a multi-agent architecture comprising five "
    "autonomous, collaborative agents: (A) a Data Ingestion Agent that cleans and engineers "
    "features from the PIMA Indians Diabetes Dataset, (B) a Risk Stratification Agent that "
    "uses an ensemble of Logistic Regression, Random Forest, and XGBoost classifiers to predict "
    "diabetes probability, (C) an Explainability Agent that leverages SHAP (SHapley Additive "
    "exPlanations) to provide transparent, clinician-friendly feature-level explanations, "
    "(D) a Recommendation Agent that uses Google Gemini via LangChain to generate personalized "
    "lifestyle and dietary recommendations, and (E) a Monitoring and Orchestrator Agent that "
    "performs longitudinal trend detection, alert generation, and autonomous LLM-based clinical "
    "reasoning over the full system state."
)
add_body(
    "The agents communicate through a shared state mechanism, enabling a closed-loop feedback "
    "system. Evaluation on the PIMA dataset yielded a best ROC-AUC of 0.84 using Random Forest. "
    "A Streamlit-based dashboard provides an interactive, real-time interface for patients and "
    "clinicians. This system demonstrates that agentic AI architectures can significantly enhance "
    "clinical decision support beyond traditional pipeline-based approaches."
)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("TABLE OF CONTENTS", 1)
toc_items = [
    ("Declaration", "ii"),
    ("Certificate", "iii"),
    ("Acknowledgement", "iv"),
    ("Abstract", "v"),
    ("Table of Contents", "vi"),
    ("List of Tables", "vii"),
    ("List of Figures", "viii"),
    ("Chapter 1: Introduction", "1"),
    ("  1.1 Background", "1"),
    ("  1.2 Problem Statement", "2"),
    ("  1.3 Limitations of Existing Systems", "3"),
    ("  1.4 Motivation", "3"),
    ("  1.5 Objectives", "4"),
    ("  1.6 Scope of the Project", "4"),
    ("Chapter 2: Literature Review", "5"),
    ("  2.1 Machine Learning in Healthcare", "5"),
    ("  2.2 Diabetes Prediction Systems", "6"),
    ("  2.3 Explainable AI and SHAP", "7"),
    ("  2.4 Large Language Models in Healthcare", "8"),
    ("  2.5 Multi-Agent Systems", "9"),
    ("  2.6 Comparison with Existing Approaches", "10"),
    ("Chapter 3: Requirement Analysis and Solution Approach", "11"),
    ("  3.1 Functional Requirements", "11"),
    ("  3.2 Non-Functional Requirements", "12"),
    ("  3.3 System Overview", "12"),
    ("  3.4 Agent Descriptions", "13"),
    ("Chapter 4: Modeling and Implementation", "17"),
    ("  4.1 Dataset Description", "17"),
    ("  4.2 Data Preprocessing", "18"),
    ("  4.3 Model Training and Selection", "20"),
    ("  4.4 Explainability with SHAP", "22"),
    ("  4.5 Agentic Architecture", "23"),
    ("  4.6 Technologies Used", "25"),
    ("Chapter 5: Testing and Results", "27"),
    ("  5.1 Model Evaluation Metrics", "27"),
    ("  5.2 Test Cases", "28"),
    ("  5.3 Sample Outputs", "30"),
    ("  5.4 Error Handling", "31"),
    ("Chapter 6: Conclusion and Future Work", "32"),
    ("  6.1 Summary", "32"),
    ("  6.2 Key Achievements", "33"),
    ("  6.3 Limitations", "33"),
    ("  6.4 Future Work", "34"),
    ("References", "35"),
    ("Ethical Considerations", "37"),
]
for item, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    indent = item.startswith("  ")
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(item.strip())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r2 = p.add_run(f"  {'.' * (50 - len(item))}  {pg}")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

print("Part 1 complete: Cover through TOC")
# Save intermediate state
doc.save(OUTPUT_PATH)
print(f"Saved to {OUTPUT_PATH}")
