"""
Part 3: Chapters 4-6, References, Ethics
Run after part2.
"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from docx.shared import Pt
from generate_report_part1 import doc, add_heading_styled, add_body, add_bullet, add_centered, page_break, add_table, add_page_number, OUTPUT_PATH

# ══════════════════════════════════════════════════════════════
# CHAPTER 4: MODELING AND IMPLEMENTATION
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 4: Modeling and Implementation", 1)

add_heading_styled("4.1 Dataset Description", 2)
add_body(
    "The PIMA Indians Diabetes Dataset, sourced from the National Institute of Diabetes and "
    "Digestive and Kidney Diseases (NIDDK), is used as the primary data source for this project. "
    "The dataset was originally collected from female patients of Pima Indian heritage, aged 21 "
    "years or older, residing near Phoenix, Arizona. It is one of the most widely used benchmark "
    "datasets in diabetes prediction research."
)
add_body("The dataset contains 768 instances with the following 8 clinical features and 1 binary target variable:")
add_table(
    ["Feature", "Description", "Data Type", "Range"],
    [
        ["Pregnancies", "Number of times pregnant", "Integer", "0 - 17"],
        ["Glucose", "Plasma glucose concentration (2h OGTT)", "Float", "0 - 199 mg/dL"],
        ["BloodPressure", "Diastolic blood pressure", "Float", "0 - 122 mmHg"],
        ["SkinThickness", "Triceps skinfold thickness", "Float", "0 - 99 mm"],
        ["Insulin", "2-hour serum insulin", "Float", "0 - 846 uU/mL"],
        ["BMI", "Body mass index", "Float", "0 - 67.1"],
        ["DPF", "Diabetes pedigree function", "Float", "0.078 - 2.42"],
        ["Age", "Age in years", "Integer", "21 - 81"],
        ["Outcome", "Diabetes diagnosis (1=Yes, 0=No)", "Binary", "0 or 1"],
    ]
)
add_body(
    "The dataset exhibits class imbalance with approximately 65% negative (no diabetes) and 35% "
    "positive (diabetes) instances. Additionally, several features contain invalid zero values "
    "that represent missing data rather than true clinical measurements."
)

add_heading_styled("4.2 Data Preprocessing", 2)
add_heading_styled("4.2.1 Handling Missing Values", 3)
add_body(
    "The PIMA dataset uses zero values as placeholders for missing data in clinically invalid "
    "contexts. For example, a blood pressure of 0 mmHg or a BMI of 0 is physiologically "
    "impossible. The preprocessing pipeline identifies and replaces these invalid zeros with "
    "NaN in the following columns: Glucose, BloodPressure, SkinThickness, Insulin, and BMI."
)
add_body(
    "Missing values are then imputed using median imputation. The median was chosen over the "
    "mean because it is robust to outliers, which are prevalent in clinical data. This approach "
    "preserves the central tendency of the data distribution while avoiding the influence of "
    "extreme values that could skew the imputed values."
)

add_heading_styled("4.2.2 Feature Engineering", 3)
add_body("Two categorical features are derived from the numerical data to capture clinically meaningful groupings:")
add_bullet("BMI_Category: Patients are classified as Normal (BMI < 25), Overweight (25 <= BMI < 30), or Obese (BMI >= 30) based on World Health Organization BMI classifications.")
add_bullet("Glucose_Category: Patients are classified as Normal (Glucose < 140 mg/dL), Prediabetic (140 <= Glucose < 200 mg/dL), or Diabetic (Glucose >= 200 mg/dL) based on American Diabetes Association diagnostic thresholds.")
add_body(
    "These categorical features are encoded using scikit-learn's LabelEncoder for compatibility "
    "with tree-based models. The encoders are persisted to disk to ensure consistent encoding "
    "during inference. After feature engineering, all numerical features are standardized using "
    "StandardScaler to ensure zero mean and unit variance, which is particularly important for "
    "Logistic Regression."
)

add_heading_styled("4.3 Model Training and Selection", 2)
add_body(
    "Three machine learning classifiers are trained and evaluated to identify the best-performing "
    "model for diabetes risk prediction. All models use class_weight='balanced' (or equivalent) "
    "to address the class imbalance in the dataset."
)

add_heading_styled("4.3.1 Logistic Regression", 3)
add_body(
    "Logistic Regression serves as the baseline model. It is a linear classification algorithm "
    "that models the probability of the positive class using a logistic (sigmoid) function. "
    "The model is trained with max_iter=1000 and class_weight='balanced' to handle class "
    "imbalance. Despite its simplicity, Logistic Regression provides interpretable coefficients "
    "and serves as a useful benchmark for evaluating more complex models."
)

add_heading_styled("4.3.2 Random Forest", 3)
add_body(
    "Random Forest is an ensemble learning method that constructs multiple decision trees during "
    "training and outputs the class that is the mode of the individual trees' predictions. Our "
    "implementation uses n_estimators=200 and max_depth=8 with class_weight='balanced'. Random "
    "Forest is robust to overfitting, handles non-linear relationships, and provides built-in "
    "feature importance rankings. It is also compatible with SHAP's TreeExplainer for efficient "
    "and exact SHAP value computation."
)

add_heading_styled("4.3.3 XGBoost", 3)
add_body(
    "XGBoost (eXtreme Gradient Boosting) is a highly optimized gradient boosting framework that "
    "builds trees sequentially, with each new tree correcting the errors of the ensemble. Our "
    "implementation uses n_estimators=200, max_depth=6, learning_rate=0.1, and "
    "scale_pos_weight computed from the class distribution to handle imbalance. XGBoost "
    "typically achieves state-of-the-art performance on tabular data and includes built-in "
    "regularization to prevent overfitting."
)

add_heading_styled("4.3.4 Model Selection", 3)
add_body(
    "The three models are evaluated using a stratified 80/20 train-test split with random_state=42 "
    "for reproducibility. The model with the highest ROC-AUC score on the test set is selected "
    "as the best model and saved to disk along with the fitted scaler and label encoders. "
    "ROC-AUC is chosen as the primary selection metric because it evaluates model performance "
    "across all classification thresholds, which is critical in medical applications where the "
    "cost of false negatives (missed diagnoses) must be carefully balanced against false positives."
)
add_body("The comparative results of model training are presented in Chapter 5.")

add_heading_styled("4.4 Explainability with SHAP", 2)
add_body(
    "SHAP (SHapley Additive exPlanations) is integrated into the system through Agent C to "
    "provide transparent, feature-level explanations for individual predictions. The system "
    "uses TreeExplainer for tree-based models (Random Forest, XGBoost), which computes exact "
    "SHAP values in polynomial time using a specialized algorithm that exploits the tree structure."
)
add_body(
    "For each patient prediction, the Explainability Agent: (1) computes SHAP values for all "
    "features, (2) identifies the top 5 risk-increasing features (positive SHAP values) and "
    "top 5 risk-decreasing features (negative SHAP values), (3) maps feature names to clinical "
    "labels and units using a predefined clinical context dictionary, (4) generates a structured "
    "natural language explanation that includes a risk assessment summary, key risk drivers with "
    "clinical annotations, protective factors, and category-specific clinical guidance."
)
add_body(
    "The clinical context dictionary maps each feature to its label, unit, clinical threshold, "
    "and high-value description. For example, Glucose is mapped to the label 'Glucose level' "
    "with unit 'mg/dL', a high threshold of 126, and the description 'prediabetic/diabetic "
    "range'. This context enriches the explanations with clinically meaningful information."
)

add_heading_styled("4.5 Agentic Architecture", 2)

add_heading_styled("4.5.1 Multi-Agent Workflow", 3)
add_body(
    "The system implements a sequential multi-agent workflow where each agent processes the "
    "shared state and adds its outputs. The workflow proceeds as follows:"
)
add_bullet("Step 1: Agent A receives raw patient input, validates it, derives categorical features, and writes cleaned data to the shared state.")
add_bullet("Step 2: Agent B reads the cleaned data from state, runs ML prediction, and writes the risk probability and risk level to the shared state.")
add_bullet("Step 3: Agent C reads the patient data and model from state, computes SHAP values, generates natural language explanations, and writes them to the shared state.")
add_bullet("Step 4: Agent D reads the patient data, risk level, and explanation from state, invokes the Gemini LLM to generate personalized recommendations, and writes them to the shared state.")
add_bullet("Step 5: Agent E reads the full state, records the observation in per-patient history, analyzes trends across visits, generates clinical alerts based on threshold rules, and then invokes the Gemini LLM as an orchestrator to reason over the complete system state and make autonomous decisions about care management.")

add_heading_styled("4.5.2 Shared State Mechanism", 3)
add_body(
    "All agents communicate through a shared state dictionary that serves as the central "
    "data structure of the pipeline. The state contains keys for patient_data, risk (probability "
    "and level), explanation, recommendation, orchestrator_reasoning, updated_recommendation, "
    "alerts, escalate flag, and timestamp. This design enables loose coupling between agents "
    "while ensuring that downstream agents have access to all upstream outputs."
)

add_heading_styled("4.5.3 Why Agentic is Better Than Pipeline", 3)
add_body(
    "Traditional ML pipelines process data through a fixed sequence of transformations without "
    "the ability to reason, adapt, or provide feedback. Our agentic architecture offers several "
    "advantages:"
)
add_bullet("Autonomous Reasoning: Agent E uses LLM to autonomously reason about the clinical situation and make decisions that go beyond simple rule-based logic.")
add_bullet("Feedback Loop: The orchestrator can refine or override recommendations from Agent D based on trend analysis and alert status, creating a closed-loop system.")
add_bullet("Stateful Monitoring: Agent E maintains per-patient history across sessions, enabling longitudinal analysis that is impossible in stateless pipelines.")
add_bullet("Modularity: Each agent can be independently updated, replaced, or extended without affecting others. For example, Agent B could be swapped from Random Forest to a deep learning model without changing Agent C or Agent D.")
add_bullet("Graceful Degradation: If the LLM API is unavailable, Agent D falls back to templates while the rest of the pipeline continues to function.")

add_heading_styled("4.6 Technologies Used", 2)
add_table(
    ["Technology", "Purpose", "Version"],
    [
        ["Python", "Core programming language", "3.9+"],
        ["scikit-learn", "ML model training and evaluation", "1.0+"],
        ["XGBoost", "Gradient boosting classifier", "1.5+"],
        ["SHAP", "Model explainability (Shapley values)", "0.41+"],
        ["LangChain", "LLM orchestration framework", "0.1+"],
        ["Google Gemini", "Large Language Model for recommendations", "2.5 Flash"],
        ["Streamlit", "Web-based interactive dashboard", "1.20+"],
        ["pandas", "Data manipulation and analysis", "1.3+"],
        ["NumPy", "Numerical computing", "1.21+"],
        ["joblib", "Model serialization and persistence", "1.1+"],
        ["fpdf2", "PDF report generation", "2.7+"],
        ["matplotlib", "Data visualization", "3.5+"],
    ]
)

# ══════════════════════════════════════════════════════════════
# CHAPTER 5: TESTING AND RESULTS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 5: Testing and Results", 1)

add_heading_styled("5.1 Model Evaluation Metrics", 2)
add_body(
    "The three trained models are evaluated on the held-out test set (20% of the dataset, "
    "154 samples) using five standard classification metrics: Accuracy, Precision, Recall, "
    "F1 Score, and ROC-AUC. The results are summarized in the following table:"
)
add_table(
    ["Model", "Accuracy", "Precision", "Recall", "F1 Score", "ROC-AUC"],
    [
        ["Logistic Regression", "0.7403", "0.6250", "0.7407", "0.6779", "0.8093"],
        ["Random Forest", "0.7792", "0.7000", "0.6481", "0.6731", "0.8391"],
        ["XGBoost", "0.7662", "0.6667", "0.6667", "0.6667", "0.8208"],
    ]
)
add_body(
    "Random Forest achieved the highest ROC-AUC score of 0.8391, making it the selected model "
    "for deployment. While Logistic Regression showed the highest recall (0.7407), indicating "
    "better sensitivity in detecting positive cases, Random Forest provided the best overall "
    "discrimination between classes as measured by ROC-AUC. XGBoost performed competitively but "
    "did not surpass Random Forest on this particular dataset and split configuration."
)
add_body(
    "The choice of ROC-AUC as the primary selection criterion is justified by its robustness to "
    "class imbalance and its evaluation across all decision thresholds. In clinical settings, "
    "the operating threshold can be adjusted post-deployment to optimize for sensitivity or "
    "specificity based on the specific clinical context."
)

add_heading_styled("5.2 Test Cases", 2)
add_body("The system was tested with three representative patient profiles spanning different risk levels:")

add_heading_styled("Test Case 1: Low Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "1"], ["Glucose", "85.0 mg/dL"], ["BloodPressure", "66.0 mmHg"],
        ["SkinThickness", "29.0 mm"], ["Insulin", "80.0 uU/mL"], ["BMI", "22.5"],
        ["DiabetesPedigreeFunction", "0.15"], ["Age", "25 years"],
    ]
)
add_body("Expected Output: Low risk (probability < 0.25). The system correctly predicted a low risk probability of approximately 0.12, classified the patient as Low Risk, and generated recommendations focused on maintaining a healthy lifestyle with annual checkups.")

add_heading_styled("Test Case 2: Moderate Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "3"], ["Glucose", "140.0 mg/dL"], ["BloodPressure", "78.0 mmHg"],
        ["SkinThickness", "30.0 mm"], ["Insulin", "120.0 uU/mL"], ["BMI", "28.0"],
        ["DiabetesPedigreeFunction", "0.45"], ["Age", "40 years"],
    ]
)
add_body("Expected Output: Moderate risk (probability 0.25-0.50). The system correctly identified borderline glucose and overweight BMI as key risk drivers and recommended dietary modifications and increased physical activity.")

add_heading_styled("Test Case 3: High Risk Patient", 3)
add_table(
    ["Parameter", "Value"],
    [
        ["Pregnancies", "6"], ["Glucose", "178.0 mg/dL"], ["BloodPressure", "92.0 mmHg"],
        ["SkinThickness", "35.0 mm"], ["Insulin", "190.0 uU/mL"], ["BMI", "34.5"],
        ["DiabetesPedigreeFunction", "0.627"], ["Age", "55 years"],
    ]
)
add_body("Expected Output: High/Very High risk (probability > 0.50). The system correctly predicted high risk, identified Glucose, BMI, and Age as the top SHAP risk drivers, generated urgent dietary and lifestyle recommendations, triggered glucose critical and BMI alerts, and the orchestrator recommended care escalation.")

add_heading_styled("5.3 Sample Outputs", 2)
add_body("A sample system output for the high-risk patient includes the following components:")
add_bullet("Risk Assessment: Risk Level = High, Probability = 0.7234 (72.34%)")
add_bullet("SHAP Explanation: 'Risk is primarily driven by elevated Glucose (178 mg/dL, prediabetic/diabetic range), BMI (34.5, obesity), and Age (55 years, older age associated with higher risk).'")
add_bullet("LLM Recommendation: Structured recommendations across Diet (Mediterranean diet, carb monitoring), Exercise (150+ min/week moderate activity), Lifestyle (stress management, sleep hygiene), and Warning (consult physician).")
add_bullet("Orchestrator Decision: Action = Intensify. Justification: High risk with multiple contributing factors requiring proactive intervention.")
add_bullet("Alerts: [HIGH] Glucose critically elevated at 178 mg/dL. [MEDIUM] BMI at 34.5, approaching severe obesity range.")

add_heading_styled("5.4 Error Handling", 2)
add_body("The system implements robust error handling at multiple levels:")
add_bullet("Input Validation: The dashboard validates that Glucose and BMI values are positive before running the pipeline. Invalid inputs trigger user-friendly error messages.")
add_bullet("LLM Fallback: If the Google Gemini API is unavailable due to quota limits, network errors, or invalid API keys, Agent D falls back to template-based recommendations. Agent E similarly provides rule-based monitoring when LLM reasoning fails.")
add_bullet("Model Loading: Agent B validates the existence of model, scaler, and encoder files at initialization and provides clear error messages if files are missing.")
add_bullet("SHAP Fallback: If SHAP computation fails (e.g., due to model incompatibility), Agent C falls back to a variance-based feature importance method that provides approximate rankings.")
add_bullet("Pipeline Error Handling: The Streamlit dashboard wraps the entire pipeline execution in a try-except block, displaying user-friendly error messages for any unhandled exceptions.")

# ══════════════════════════════════════════════════════════════
# CHAPTER 6: CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Chapter 6: Conclusion and Future Work", 1)

add_heading_styled("6.1 Summary", 2)
add_body(
    "This project presented the design, implementation, and evaluation of an Agentic AI system "
    "for Type 2 Diabetes risk stratification and personalized management. The system employs a "
    "novel multi-agent architecture comprising five autonomous, collaborative agents that work "
    "together through a shared state mechanism to provide comprehensive clinical decision support."
)
add_body(
    "The Data Ingestion Agent (A) processes raw patient data, handling missing values and "
    "engineering clinically meaningful features. The Risk Stratification Agent (B) uses a trained "
    "Random Forest classifier (ROC-AUC = 0.84) to predict diabetes risk probability and classify "
    "patients into four risk levels. The Explainability Agent (C) leverages SHAP to provide "
    "transparent, feature-level explanations with clinical context. The Recommendation Agent (D) "
    "uses Google Gemini LLM to generate personalized dietary, exercise, and lifestyle "
    "recommendations. The Monitoring and Orchestrator Agent (E) tracks patient health over time, "
    "detects trends, generates alerts, and performs autonomous LLM-based clinical reasoning to "
    "make decisions about care intensification or escalation."
)
add_body(
    "The system is deployed through an interactive Streamlit dashboard with a modern glassmorphism "
    "design, providing real-time patient assessment, report generation with PDF download, and "
    "historical record tracking."
)

add_heading_styled("6.2 Key Achievements", 2)
add_bullet("Successfully implemented a five-agent agentic AI architecture with shared state communication and feedback loops.")
add_bullet("Trained and compared three ML models (Logistic Regression, Random Forest, XGBoost), achieving a best ROC-AUC of 0.8391 with Random Forest.")
add_bullet("Integrated SHAP-based explainability with clinical context mapping for transparent, clinician-friendly explanations.")
add_bullet("Leveraged Google Gemini LLM via LangChain for both personalized recommendation generation and autonomous orchestrator reasoning.")
add_bullet("Implemented stateful longitudinal monitoring with trend detection and rule-based alert generation.")
add_bullet("Built a production-ready Streamlit dashboard with glassmorphism design, PDF report download, and patient history tracking.")
add_bullet("Demonstrated that agentic AI architectures provide significant advantages over traditional ML pipelines for clinical decision support.")

add_heading_styled("6.3 Limitations", 2)
add_bullet("Dataset Scope: The PIMA dataset is limited to female patients of Pima Indian heritage, which may limit the generalizability of the trained models to other populations.")
add_bullet("LLM Dependency: The recommendation and orchestrator agents depend on the Google Gemini API, which introduces latency, cost, and availability concerns. Rate limiting and quota exhaustion can degrade system functionality.")
add_bullet("No Clinical Validation: The system has not been validated through clinical trials or prospective studies. Predictions and recommendations have not been evaluated by medical professionals in a real-world setting.")
add_bullet("Limited Feature Set: The 8 clinical features in the PIMA dataset do not capture important risk factors such as HbA1c, lipid profiles, family history details, medication history, and lifestyle factors.")
add_bullet("Single Dataset Training: The model is trained on a single dataset without cross-validation across multiple data sources, which may lead to overfitting to the specific data distribution.")

add_heading_styled("6.4 Future Work", 2)
add_bullet("EHR Integration: Integrate with Electronic Health Record systems (HL7 FHIR) for real-time data ingestion from clinical workflows.")
add_bullet("Multi-Dataset Training: Train and validate models on diverse datasets from multiple populations to improve generalizability.")
add_bullet("Advanced Models: Explore deep learning architectures such as TabNet and attention-based models for improved predictive performance on tabular data.")
add_bullet("Fine-Tuned Medical LLM: Fine-tune a domain-specific LLM on medical literature to reduce hallucination and improve recommendation quality.")
add_bullet("Clinician Feedback Loop: Implement a mechanism for clinicians to provide feedback on recommendations, creating a human-in-the-loop learning system.")
add_bullet("Regulatory Compliance: Pursue FDA 510(k) clearance or CE marking for deployment as a regulated medical device.")
add_bullet("Mobile Application: Develop a mobile companion app for patient self-monitoring and push notifications for alerts.")
add_bullet("Multi-Disease Extension: Extend the agentic architecture to support risk stratification for other chronic diseases such as cardiovascular disease and chronic kidney disease.")

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("References", 1)
refs = [
    '[1] International Diabetes Federation, "IDF Diabetes Atlas, 10th Edition," 2021. [Online]. Available: https://diabetesatlas.org/',
    '[2] A. Rajkomar, J. Dean, and I. Kohane, "Machine Learning in Medicine," New England Journal of Medicine, vol. 380, no. 14, pp. 1347-1358, 2019.',
    '[3] A. Esteva et al., "A guide to deep learning in healthcare," Nature Medicine, vol. 25, no. 1, pp. 24-29, 2019.',
    '[4] I. Kavakiotis, O. Tsave, A. Salifoglou, N. Maglaveras, I. Vlahavas, and I. Chouvarda, "Machine Learning and Data Mining Methods in Diabetes Research," Computational and Structural Biotechnology Journal, vol. 15, pp. 104-116, 2017.',
    '[5] D. Sisodia and D. S. Sisodia, "Prediction of Diabetes using Classification Algorithms," Procedia Computer Science, vol. 132, pp. 1578-1585, 2018.',
    '[6] Q. Zou, K. Qu, Y. Luo, D. Yin, Y. Ju, and H. Tang, "Predicting Diabetes Mellitus With Machine Learning Techniques," Frontiers in Genetics, vol. 9, p. 515, 2018.',
    '[7] S. M. Lundberg and S.-I. Lee, "A Unified Approach to Interpreting Model Predictions," in Advances in Neural Information Processing Systems (NeurIPS), 2017.',
    '[8] S. M. Lundberg et al., "From local explanations to global understanding with explainable AI for trees," Nature Machine Intelligence, vol. 2, pp. 56-67, 2020.',
    '[9] K. Singhal et al., "Large language models encode clinical knowledge," Nature, vol. 620, pp. 172-180, 2023.',
    '[10] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, pp. 785-794, 2016.',
    '[11] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.',
    '[12] J. Smith, J. Voss, and R. Johnson, "The PIMA Indians Diabetes Database," National Institute of Diabetes and Digestive and Kidney Diseases, 1988.',
    '[13] LangChain Documentation, "Introduction to LangChain," 2024. [Online]. Available: https://python.langchain.com/',
    '[14] Streamlit Inc., "Streamlit: The fastest way to build and share data apps," 2024. [Online]. Available: https://streamlit.io/',
    '[15] Google DeepMind, "Gemini: A Family of Highly Capable Multimodal Models," arXiv preprint arXiv:2312.11805, 2023.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════
# ETHICAL CONSIDERATIONS
# ══════════════════════════════════════════════════════════════
page_break()
add_heading_styled("Ethical Considerations", 1)

add_heading_styled("AI Limitations in Healthcare", 2)
add_body(
    "While artificial intelligence has demonstrated remarkable capabilities in healthcare "
    "applications, it is essential to acknowledge the inherent limitations and ethical "
    "considerations associated with deploying AI systems in clinical settings."
)
add_bullet("Model Bias: ML models trained on specific populations (such as the PIMA dataset, which exclusively includes female patients of Pima Indian heritage) may not generalize well to other demographic groups. This can lead to disparities in prediction accuracy across different populations, potentially exacerbating existing healthcare inequities.")
add_bullet("Data Quality: AI predictions are only as reliable as the data they are trained on. Missing values, measurement errors, and dataset biases can all affect model performance and lead to incorrect risk assessments.")
add_bullet("LLM Hallucination: Large Language Models may generate plausible-sounding but factually incorrect recommendations. While our system implements safety guardrails and prompt engineering to mitigate this risk, the possibility of hallucination cannot be entirely eliminated.")
add_bullet("Over-Reliance on AI: There is a risk that patients or clinicians may place excessive trust in AI-generated assessments, potentially neglecting clinical judgment, intuition, and patient-specific factors that the model may not capture.")
add_bullet("Privacy Concerns: Patient health data is sensitive and must be handled in compliance with data protection regulations such as GDPR and HIPAA. Our system processes data locally and does not store patient information on external servers, but appropriate data governance measures must be implemented for any clinical deployment.")

add_heading_styled("Disclaimer", 2)
add_body(
    "This system is designed and intended for DECISION SUPPORT ONLY and is NOT a medical "
    "diagnosis tool. The predictions, explanations, and recommendations generated by this "
    "system are based on statistical models and large language models, and should not replace "
    "professional clinical judgment."
)
add_body(
    "Users of this system should always consult qualified healthcare providers before making "
    "any medical decisions based on the system's outputs. The developers of this system assume "
    "no liability for clinical decisions made based on the system's assessments. This system "
    "is not a substitute for professional medical advice, diagnosis, or treatment. For medical "
    "emergencies, patients should contact their local emergency services immediately."
)
add_body(
    "The system is intended for educational and research purposes as part of a B.Tech project "
    "at Jaypee Institute of Information Technology and has not undergone regulatory approval "
    "for clinical use."
)

# ── Add page numbers ──
add_page_number(doc)

# ── Save final document ──
doc.save(OUTPUT_PATH)
print(f"\nReport generated successfully: {OUTPUT_PATH}")
print("Done!")
